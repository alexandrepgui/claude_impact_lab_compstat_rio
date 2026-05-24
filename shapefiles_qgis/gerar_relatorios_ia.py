# -*- coding: utf-8 -*-
"""
Gera RELINTs (.docx) por zona usando Claude Sonnet 4.6.

Para cada zona da última semana (limitado por --max-zonas):
  1) Lê o dossiê compacto (distribuicao_fm/relatorio_zonas_compacto.json)
  2) Renderiza um mapa Leaflet → screenshot PNG via Edge headless
  3) Chama Claude com (dossiê + RELINT exemplar) e gera texto no estilo dos
     documentos em relints/
  4) Monta .docx com texto + mapa

Saídas em distribuicao_fm/relatorios_ia/:
  - RA_<prioridade>_<local>.docx
  - RA_<prioridade>_<local>_mapa.png

Requisitos:
  pip install anthropic python-docx
  set ANTHROPIC_API_KEY=sk-ant-...
  Edge instalado (caminho hard-coded abaixo)
"""
from __future__ import annotations
import argparse
import json
import math
import os
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

import shapefile  # pyshp
import anthropic
from docx import Document
from docx.shared import Cm, Pt, RGBColor

BASE = Path(__file__).resolve().parent
FM = BASE / "distribuicao_fm"
RELATORIO_JSON = FM / "relatorio_zonas_compacto.json"
ZONAS_SHP = FM / "zonas_semanais"
OUT_DIR = FM / "relatorios_ia"

# Renderização do mapa usa Playwright (cross-platform). Instalar:
#   pip install playwright && python -m playwright install chromium
MODEL = "claude-sonnet-4-6"

# Camadas de pontos pra plotar no mapa (latin-1 universal — vide relatorio_zonas.py)
# (nome_logico, path, cor, max_amostra_no_mapa)
POINT_LAYERS = [
    ("ocorrencias",     "ocorrencias/ocorrencias",         "#d62728", 1000),
    ("disque",          "disk_denuncia/disk_denuncia",     "#ff7f0e", 400),
    ("fatores_urbanos", "fatores_urbanos/fatores_urbanos", "#1f77b4", 400),
]


# RELINT exemplar — extraído de relints/RI_017 (Presidente Vargas).
# Cache_control marker no system prompt → usado em prompt caching.
RELINT_EXEMPLAR = """RELATÓRIO DE INTELIGÊNCIA DE ÁREA – COMPSTAT – DADOS PÚBLICOS
RELATÓRIO DE INTELIGÊNCIA DE ÁREA
Subsídio para Reunião de CompStat
PRESIDENTE VARGAS – CAMPO DE SANTANA – CENTRAL – CINELÂNDIA

A presente análise territorial visa identificar fatores urbanos, dinâmica criminal e vulnerabilidades relacionadas à sensação de segurança da população, considerando áreas de grande circulação de pedestres e integração com o transporte público. Foram observados fatores associados à incidência de furtos e roubos contra transeuntes, especialmente subtração de aparelhos celulares, além de elementos urbanos que favorecem delitos oportunistas.

AVENIDA PRESIDENTE VARGAS
A Avenida Presidente Vargas concentra intenso fluxo de pedestres e usuários do transporte público, favorecendo furtos e roubos de celulares em pontos de ônibus e calçadas. Com extensão de 4,5 km entre a Praça da República e o Cais do Porto, a via é a principal artéria do Centro e opera como rota de escape para criminosos após ações nas ruas adjacentes. O cruzamento com a Rua Uruguaiana funciona como ponto de escoamento de produtos subtraídos para o Camelódromo da Uruguaiana, maior polo de receptação de celulares furtados do estado.

Também foram identificados:
- retenção de fluxo em horários de pico — pontos de ônibus ao longo de toda a extensão da via;
- áreas com baixa visibilidade — trechos sob marquises, estruturas elevadas e passarelas;
- obstáculos urbanos dificultando vigilância — ambulantes irregulares obstruindo calçadas e comprometendo visibilidade;
- circulação intensa de motocicletas e bicicletas — incluindo motocicletas com placas adulteradas identificadas no corredor;
- múltiplas rotas de dispersão após a prática criminosa — vias transversais a cada quadra viabilizam fuga imediata.

A dinâmica criminal observada indica predominância de indivíduos atuando a pé, motocicletas e bicicletas, aproveitando momentos de distração das vítimas em áreas de espera, travessias e acessos ao transporte público.

CAMPO DE SANTANA
O Campo de Santana (Praça da República), tombado pelo INEPAC, apresenta vegetação densa e áreas abertas utilizadas para dispersão após delitos praticados nas vias adjacentes. Vistoria realizada constatou presença de população em situação de rua em múltiplos pontos do parque, usuários de entorpecentes com abordagens agressivas a transeuntes, lixo acumulado e falha grave na iluminação de setores internos. Os pontos de ônibus na Av. Marechal Floriano (fronteira sul) e na Av. Presidente Vargas (fronteira norte) apresentam grande concentração de usuários que cruzam o parque diariamente, tornando o trajeto de risco permanente.

Também foram identificados:
- retenção de fluxo em horários de pico — pedestres cruzando o parque entre os terminais de ônibus adjacentes;
- áreas com baixa visibilidade — setores internos sem iluminação e vegetação densa favorecendo abordagens noturnas;
- obstáculos urbanos dificultando vigilância — bancos e estruturas de lazer ocupados por PSR como pontos fixos;
- circulação intensa de motocicletas e bicicletas — vias do entorno utilizadas como corredores de fuga após delitos;
- múltiplas rotas de dispersão após a prática criminosa — saídas abertas para a Av. Presidente Vargas e Av. Marechal Floriano.

CONCLUSÃO
A área analisada apresenta fatores urbanos e territoriais que favorecem delitos oportunistas, especialmente furtos e roubos de aparelhos celulares contra pedestres. A Av. Presidente Vargas opera como corredor de escoamento de produtos subtraídos; o Campo de Santana funciona como ponto de apoio e dispersão; e a Central do Brasil concentra o maior fluxo de passageiros do país em um nó de múltiplas rotas de fuga. Observa-se necessidade de:

- reforço do patrulhamento preventivo — com missões dirigidas por horário e local, priorizando os acessos à Central do Brasil (07h–09h e 17h–20h) e a Praça Floriano no período noturno;
- melhoria da iluminação pública — especialmente nos setores internos do Campo de Santana e no entorno da Praça Floriano;
- poda de vegetação — redução dos pontos de sombra no Campo de Santana que favorecem abordagens e permanência de PSR;
- fiscalização e ordenamento do comércio irregular — ações específicas no Camelódromo da Uruguaiana e nos acessos à Central do Brasil;
- ações integradas de ordenamento urbano — coordenação permanente entre Força Municipal, PMERJ, PCERJ e SEOP, com operações abrangendo o eixo completo do perímetro.

Os delitos tendem a ocorrer principalmente nos horários de pico de transporte público (07h–09h e 17h–20h) na Central do Brasil e Av. Presidente Vargas, e no período noturno na Cinelândia, com agravamento nos fins de semana e períodos de grande aglomeração.
"""


SYSTEM_PROMPT = f"""Você é um analista do CompStat Rio escrevendo um Relatório
de Inteligência de Área (RELINT) com base em dados de uma zona prioritária da
Força Municipal. O relatório serve de subsídio para a reunião semanal de
CompStat e segue um formato específico, ilustrado pelo exemplo abaixo.

Diretrizes:
- Use linguagem técnica e impessoal, sem floreios, no estilo do exemplo.
- Estruture o documento como: título "RELATÓRIO DE INTELIGÊNCIA DE ÁREA" /
  subtítulo da zona / parágrafo introdutório / uma seção por sub-localidade
  relevante (rua/bairro/landmark dominante) / CONCLUSÃO com recomendações.
- Em cada sub-seção, inclua a lista padrão de fatores urbanos quando
  aplicável (retenção de fluxo / baixa visibilidade / obstáculos /
  motos-bicicletas / rotas de dispersão), adaptada ao que os dados mostram.
- Cite números reais do dossiê (totais, horários de pico, tipos predominantes,
  bairros). NÃO invente endereços, nomes próprios, datas específicas ou
  estatísticas não presentes no dossiê.
- Quando os relatos do Disque Denúncia sugerem modus operandi ou dinâmica
  específica, parafraseie — não cite literalmente.
- Termine com a CONCLUSÃO contendo bullets de recomendação por órgão
  responsável (Força Municipal, COMLURB, SECONSERVA, Rio Luz, SMAS, SEOP,
  CET-Rio, GM-Rio, conforme os fatores urbanos da zona indicarem).
- Tamanho: ~600 a 900 palavras.

EXEMPLO de RELINT (estilo a replicar):

{RELINT_EXEMPLAR}
"""


# --------------------------------------------------------------------------- #
# Geometria e helpers
# --------------------------------------------------------------------------- #
def fix_mojibake(s):
    if not isinstance(s, str):
        return s
    try:
        return s.encode("latin-1").decode("utf-8")
    except (UnicodeEncodeError, UnicodeDecodeError):
        return s


def carregar_zona_polygon(zona_id, iso_ano, iso_sem):
    r = shapefile.Reader(str(ZONAS_SHP), encoding="utf-8")
    for s, rec in zip(r.iterShapes(), r.iterRecords()):
        if (rec["zona_id"] == zona_id
                and rec["iso_ano"] == iso_ano
                and rec["iso_sem"] == iso_sem):
            return list(s.points)
    raise ValueError(f"Zona {zona_id} de {iso_ano}-S{iso_sem:02d} não achada.")


def pontos_no_bbox(layer_path, bbox, encoding="latin-1", max_amostra=None,
                   seed=0):
    """Devolve [(lon,lat), ...] da camada que caem dentro do bbox.
    Se max_amostra for fornecido e a camada exceder, faz amostra aleatória
    determinística (seed = zona_id) para não embolar o mapa.
    """
    import random
    x0, y0, x1, y1 = bbox
    r = shapefile.Reader(str(BASE / layer_path), encoding=encoding)
    pts = []
    for s in r.iterShapes():
        if not s.points:
            continue
        lon, lat = s.points[0]
        if x0 <= lon <= x1 and y0 <= lat <= y1:
            pts.append((round(lon, 6), round(lat, 6)))
    if max_amostra and len(pts) > max_amostra:
        rng = random.Random(seed)
        pts = rng.sample(pts, max_amostra)
    return pts


def expandir_bbox(bbox, fator=0.25):
    x0, y0, x1, y1 = bbox
    dx, dy = (x1 - x0) * fator, (y1 - y0) * fator
    return (x0 - dx, y0 - dy, x1 + dx, y1 + dy)


# --------------------------------------------------------------------------- #
# Renderização do mapa
# --------------------------------------------------------------------------- #
HTML_MAPA = r"""<!DOCTYPE html><html><head><meta charset="utf-8"/>
<link rel="stylesheet" href="https://unpkg.com/leaflet@1.9.4/dist/leaflet.css"/>
<script src="https://unpkg.com/leaflet@1.9.4/dist/leaflet.js"></script>
<style>html,body,#map{margin:0;height:100%;}
.lbl{background:rgba(0,0,0,.7);color:#fff;border-radius:4px;padding:2px 6px;font:bold 12px Arial;white-space:nowrap;}
.legenda{position:absolute;z-index:1000;right:8px;bottom:8px;background:rgba(255,255,255,.95);
padding:8px 10px;border-radius:6px;font:12px Arial;box-shadow:0 1px 4px rgba(0,0,0,.3)}
.dot{display:inline-block;width:10px;height:10px;border-radius:5px;margin-right:5px;vertical-align:middle}
.titulo{position:absolute;z-index:1000;left:8px;top:8px;background:rgba(255,255,255,.95);
padding:6px 10px;border-radius:6px;font:bold 14px Arial;box-shadow:0 1px 4px rgba(0,0,0,.3)}
</style></head><body>
<div id="map"></div>
<div class="titulo">__TITULO__</div>
<div class="legenda">
  <div><span class="dot" style="background:#d62728"></span>Ocorrências (__N_OC__)</div>
  <div><span class="dot" style="background:#ff7f0e"></span>Disque (__N_DQ__)</div>
  <div><span class="dot" style="background:#1f77b4"></span>Fatores urbanos (__N_FT__)</div>
  <div><span class="dot" style="background:#700;border:1px solid #700"></span>Zona FM</div>
</div>
<script>
const D = __DATA__;
const map = L.map('map', {zoomControl:true});
L.tileLayer('https://{s}.tile.openstreetmap.org/{z}/{x}/{y}.png',
  {maxZoom:19, attribution:'© OpenStreetMap'}).addTo(map);

// pontos primeiro
const COR = {ocorrencias:'#d62728', disque:'#ff7f0e', fatores_urbanos:'#1f77b4'};
for (const [nome, pts] of Object.entries(D.pontos)) {
  for (const [lon, lat] of pts) {
    L.circleMarker([lat, lon], {radius:2.5, color:COR[nome], weight:0,
      fillColor:COR[nome], fillOpacity:.55}).addTo(map);
  }
}
// polígono por cima (linha grossa + fill com leve transparência pra contraste)
const poly = L.polygon(D.poly, {color:'#400', weight:4, fillColor:'#900', fillOpacity:.07}).addTo(map);
L.marker(poly.getBounds().getCenter(),{icon:L.divIcon({className:'',html:`<span class="lbl">${D.label}</span>`,iconSize:[0,0]})}).addTo(map);

map.fitBounds(poly.getBounds().pad(0.30));
setTimeout(()=>window.__ready=true, 500);
</script></body></html>
"""


def renderizar_mapa(out_png, titulo, ring_lonlat, pontos_layers, n_oc, n_dq, n_ft):
    """Gera o HTML e dispara o Edge headless pra screenshot."""
    data = {
        "label": titulo,
        # Leaflet espera [lat, lon]
        "poly": [[round(y, 6), round(x, 6)] for x, y in ring_lonlat],
        "pontos": pontos_layers,
    }
    html = (HTML_MAPA
            .replace("__DATA__", json.dumps(data, separators=(",", ":")))
            .replace("__TITULO__", titulo)
            .replace("__N_OC__", str(n_oc))
            .replace("__N_DQ__", str(n_dq))
            .replace("__N_FT__", str(n_ft)))
    tmp_dir = tempfile.mkdtemp(prefix="relint_mapa_")
    try:
        tmp_html = Path(tmp_dir) / "mapa.html"
        tmp_html.write_text(html, encoding="utf-8")

        # Playwright headless: cross-platform, renderiza Leaflet + tiles + estilos.
        try:
            from playwright.sync_api import sync_playwright  # type: ignore
        except ImportError as e:
            raise RuntimeError(
                "Playwright não instalado. Rode: pip install -r requirements.txt "
                "&& python -m playwright install chromium"
            ) from e

        with sync_playwright() as p:
            browser = p.chromium.launch(args=["--no-sandbox"])
            page = browser.new_page(viewport={"width": 1400, "height": 1000})
            page.goto(f"file://{tmp_html}")
            try:
                page.wait_for_function("window.__ready === true", timeout=10000)
            except Exception:
                # Mapa pode ter falhado no JS — segue com screenshot do que tem
                page.wait_for_timeout(2000)
            page.screenshot(path=out_png, full_page=False)
            browser.close()

        if not Path(out_png).exists():
            raise RuntimeError("Playwright não gerou screenshot.")
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)


# --------------------------------------------------------------------------- #
# LLM
# --------------------------------------------------------------------------- #
def gerar_texto_relint(client, dossie, semana):
    user_msg = (
        f"Gere o RELINT para a zona abaixo, da semana {semana['iso_ano']}-S"
        f"{semana['iso_sem']:02d} (segunda-feira: {semana['segunda_feira']}).\n\n"
        f"DOSSIÊ DA ZONA (JSON):\n```json\n"
        f"{json.dumps(dossie, ensure_ascii=False, indent=2)}\n```\n\n"
        "Escreva o documento completo (~600-900 palavras) seguindo a estrutura "
        "do exemplo. Use os números reais do JSON acima e não invente "
        "endereços, datas ou estatísticas não presentes nele."
    )
    response = client.messages.create(
        model=MODEL,
        max_tokens=16000,
        system=[
            {"type": "text", "text": SYSTEM_PROMPT,
             "cache_control": {"type": "ephemeral"}}
        ],
        messages=[{"role": "user", "content": user_msg}],
    )
    texto = "\n".join(b.text for b in response.content if b.type == "text")
    return texto, response.usage


# --------------------------------------------------------------------------- #
# DOCX
# --------------------------------------------------------------------------- #
def montar_docx(texto, mapa_png, dossie, semana, out_docx):
    doc = Document()
    # margens
    for sec in doc.sections:
        sec.top_margin = Cm(1.8)
        sec.bottom_margin = Cm(1.8)
        sec.left_margin = Cm(2.0)
        sec.right_margin = Cm(2.0)

    # metadado curtinho no topo
    p = doc.add_paragraph()
    r = p.add_run(
        f"Semana {semana['iso_ano']}-S{semana['iso_sem']:02d} · "
        f"Segunda-feira: {semana['segunda_feira']} · "
        f"Zona #{dossie['prioridade']} ({dossie['local_rotulo']}) · "
        f"{dossie['agentes_alocados']} agentes alocados · "
        f"{dossie['n_celulas']} células · {dossie['area_km2']} km²"
    )
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # mapa
    if mapa_png and Path(mapa_png).exists():
        doc.add_picture(str(mapa_png), width=Cm(16))
        cap = doc.add_paragraph()
        cr = cap.add_run(
            f"Mapa: polígono da zona com ocorrências (vermelho), denúncias do "
            f"Disque (laranja) e fatores urbanos (azul) no entorno. "
            f"Centróide: {dossie['centroide']['lat']}, "
            f"{dossie['centroide']['lon']}."
        )
        cr.font.size = Pt(9)
        cr.italic = True
        cr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # corpo: parsea o texto retornado pelo LLM em parágrafos / títulos / bullets
    for raw_line in texto.split("\n"):
        line = raw_line.rstrip()
        if not line:
            doc.add_paragraph()
            continue
        # heurística simples: linhas em UPPER ou começando com "CONCLUSÃO"/
        # "RELATÓRIO" viram títulos
        stripped = line.strip()
        if stripped == stripped.upper() and len(stripped) > 5 and not stripped.startswith(("-", "•", "*")):
            h = doc.add_paragraph()
            run = h.add_run(stripped)
            run.bold = True
            run.font.size = Pt(12)
            continue
        if stripped.startswith(("- ", "• ", "* ")):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
            continue
        doc.add_paragraph(stripped)

    doc.save(str(out_docx))


# --------------------------------------------------------------------------- #
# Main
# --------------------------------------------------------------------------- #
def slug(s):
    return "".join(c if c.isalnum() else "_" for c in s).strip("_") or "zona"


def main():
    # PIPELINE_REPORT_MAX_ZONAS env: default 8 (todas), 0 = todas, N = primeiras N
    env_max = int(os.environ.get("PIPELINE_REPORT_MAX_ZONAS", "8"))
    parser = argparse.ArgumentParser()
    parser.add_argument("--max-zonas", type=int, default=env_max,
                        help=f"Quantas zonas processar (default {env_max}; "
                             f"env PIPELINE_REPORT_MAX_ZONAS; 0 = todas).")
    parser.add_argument("--zonas", type=str, default=None,
                        help="IDs específicos separados por vírgula (ex.: 1,3).")
    args = parser.parse_args()

    # Carrega .env via _llm_client (handles fallback se python-dotenv não estiver)
    try:
        sys.path.insert(0, str(BASE.parent))
        from pipeline_steps._llm_client import get_client, is_configured  # type: ignore
        if not is_configured():
            sys.exit("ERRO: ANTHROPIC_API_KEY não configurada. cp .env.example .env")
        client = get_client()
    except ImportError:
        if not os.environ.get("ANTHROPIC_API_KEY"):
            sys.exit("ERRO: ANTHROPIC_API_KEY não configurada.")
        client = anthropic.Anthropic()

    if not RELATORIO_JSON.exists():
        sys.exit(f"ERRO: rode relatorio_zonas.py antes; falta {RELATORIO_JSON}")

    payload = json.loads(RELATORIO_JSON.read_text(encoding="utf-8"))
    semana = payload["semana"]
    zonas = payload["zonas"]
    if args.zonas:
        alvo = set(int(x) for x in args.zonas.split(","))
        zonas = [z for z in zonas if z["zona_id"] in alvo]
    elif args.max_zonas > 0:
        zonas = zonas[:args.max_zonas]
    # se max_zonas == 0, processa todas

    OUT_DIR.mkdir(parents=True, exist_ok=True)

    print(f"Modelo: {MODEL}")
    print(f"Processando {len(zonas)} zona(s) da semana "
          f"{semana['iso_ano']}-S{semana['iso_sem']:02d}")

    for z in zonas:
        zid = z["zona_id"]
        nome = f"RA_{z['prioridade']:03d}_{slug(z['local_rotulo'])}"
        print(f"\n[Zona #{z['prioridade']} · {z['local_rotulo']}]")

        # 1) mapa
        ring = carregar_zona_polygon(zid, semana["iso_ano"], semana["iso_sem"])
        bbox_zona = (z["bbox"]["lon_min"], z["bbox"]["lat_min"],
                     z["bbox"]["lon_max"], z["bbox"]["lat_max"])
        bbox_alargado = expandir_bbox(bbox_zona, fator=0.30)
        pontos = {nome_layer: pontos_no_bbox(path, bbox_alargado,
                                              max_amostra=cap, seed=zid)
                  for nome_layer, path, _, cap in POINT_LAYERS}
        mapa_png = OUT_DIR / f"{nome}_mapa.png"
        print(f"  mapa → {mapa_png.name} "
              f"(ocor={len(pontos['ocorrencias'])}, "
              f"disque={len(pontos['disque'])}, "
              f"fatores={len(pontos['fatores_urbanos'])})")
        renderizar_mapa(
            str(mapa_png),
            titulo=f"{z['local_rotulo']} · {z['agentes_alocados']} agentes",
            ring_lonlat=ring,
            pontos_layers=pontos,
            n_oc=len(pontos["ocorrencias"]),
            n_dq=len(pontos["disque"]),
            n_ft=len(pontos["fatores_urbanos"]),
        )

        # 2) LLM
        print(f"  chamando {MODEL}...")
        texto, usage = gerar_texto_relint(client, z, semana)
        cache_read = getattr(usage, "cache_read_input_tokens", 0) or 0
        cache_write = getattr(usage, "cache_creation_input_tokens", 0) or 0
        print(f"    tokens: in={usage.input_tokens} out={usage.output_tokens} "
              f"cache_w={cache_write} cache_r={cache_read}")

        # 3) DOCX
        out_docx = OUT_DIR / f"{nome}.docx"
        montar_docx(texto, mapa_png, z, semana, out_docx)
        print(f"  docx → {out_docx.name} "
              f"({out_docx.stat().st_size/1024:.1f} KB)")


if __name__ == "__main__":
    main()
