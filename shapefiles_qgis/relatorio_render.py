# -*- coding: utf-8 -*-
"""
Renderização do RELINT em dois formatos:

  - PDF estilo "sóbrio corporativo" (ReportLab)
  - DOCX enxuto e editável (python-docx, com Title/Heading 1/Heading 2)

A entrada é o texto bruto que o LLM (Claude) gera em `gerar_relatorios_ia.py`,
que vem misturado com markdown (`#`, `##`, `**bold**`, `---`). O parser
limpa isso e produz uma IR (Documento → Blocos → Spans) que ambos os
renderers consomem.

Uso:
    from shapefiles_qgis.relatorio_render import (
        parse_relint, render_pdf, render_docx
    )

    doc_ir = parse_relint(texto_llm)
    render_pdf(doc_ir, mapa_png, dossie, semana, out_pdf)
    render_docx(doc_ir, mapa_png, dossie, semana, out_docx)
"""
from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Literal

# --------------------------------------------------------------------------- #
# IR
# --------------------------------------------------------------------------- #
@dataclass
class Span:
    text: str
    bold: bool = False
    italic: bool = False


@dataclass
class Bloco:
    tipo: Literal["heading", "paragrafo", "bullet", "lista_intro"]
    spans: list[Span] = field(default_factory=list)
    nivel: int = 1  # heading: 1 = section, 2 = subsection


@dataclass
class Documento:
    titulo: str = ""
    subsidio: str = ""
    subtitulo_zona: str = ""
    blocos: list[Bloco] = field(default_factory=list)


# --------------------------------------------------------------------------- #
# Parser
# --------------------------------------------------------------------------- #
_BOLD_PATTERN = re.compile(r"\*\*(.+?)\*\*")


def _parse_inline(texto: str) -> list[Span]:
    """Quebra `... **bold** ... *italic* ...` em spans."""
    spans: list[Span] = []
    pos = 0
    for m in _BOLD_PATTERN.finditer(texto):
        if m.start() > pos:
            spans.append(Span(texto[pos:m.start()]))
        spans.append(Span(m.group(1), bold=True))
        pos = m.end()
    if pos < len(texto):
        spans.append(Span(texto[pos:]))
    if not spans:
        spans = [Span(texto)]
    return spans


def _strip_md_heading(line: str) -> tuple[int, str]:
    """`# foo` → (1, 'foo'); `## foo` → (2, 'foo'); else (0, line)."""
    m = re.match(r"^(#{1,6})\s+(.*)$", line)
    if m:
        return len(m.group(1)), m.group(2).strip()
    return 0, line


def _strip_bullet(line: str) -> str | None:
    """`- foo` / `* foo` / `• foo` → 'foo'; else None."""
    m = re.match(r"^[-*•]\s+(.+)$", line)
    return m.group(1).strip() if m else None


def _is_separator(line: str) -> bool:
    return bool(re.match(r"^[-=_]{3,}$", line.strip()))


def _is_only_bold(line: str) -> str | None:
    """`**foo**` (linha inteira só negrito) → 'foo'; else None."""
    m = re.match(r"^\s*\*\*(.+?)\*\*\s*$", line)
    return m.group(1).strip() if m else None


def parse_relint(texto: str) -> Documento:
    """Converte o texto bruto do LLM (com markdown) em IR."""
    doc = Documento()
    # Primeiro passo: identifica cabeçalho do documento.
    linhas = [l.rstrip() for l in texto.splitlines()]
    # Remove linhas em branco do começo
    while linhas and not linhas[0].strip():
        linhas.pop(0)

    # Cabeçalho: enquanto vemos headings ou linhas só-negrito (subsídio/subtítulo),
    # acumula. Para no primeiro `---` ou no primeiro parágrafo "normal".
    cabecalho: list[str] = []
    while linhas:
        l = linhas[0].strip()
        if not l:
            linhas.pop(0)
            continue
        if _is_separator(l):
            linhas.pop(0)
            break
        nivel, _ = _strip_md_heading(l)
        bold_only = _is_only_bold(l)
        if nivel or bold_only is not None:
            cabecalho.append(l)
            linhas.pop(0)
            continue
        # parágrafo normal -> fim do cabeçalho
        break

    # Processa cabeçalho. O primeiro heading é o título; demais são tratados
    # como ruído (a LLM costuma emitir `# X` e `## X` redundantes). Subsídio
    # e subtítulo da zona vêm das linhas só-negrito (`**...**`).
    for c in cabecalho:
        c = c.strip()
        if not c:
            continue
        nivel, txt = _strip_md_heading(c)
        if nivel:
            if not doc.titulo:
                doc.titulo = txt
            continue  # ignora headings duplicados no cabeçalho
        bold_only = _is_only_bold(c)
        if bold_only:
            if not doc.subsidio:
                doc.subsidio = bold_only
            elif not doc.subtitulo_zona:
                doc.subtitulo_zona = bold_only
            else:
                doc.subtitulo_zona += " · " + bold_only

    # Corpo: parse linha-a-linha em blocos.
    paragrafo_buffer: list[str] = []

    def flush_paragrafo():
        if paragrafo_buffer:
            txt = " ".join(paragrafo_buffer).strip()
            if txt:
                # Detecta "Também foram identificados:" como lista_intro
                tipo = "lista_intro" if txt.endswith(":") and len(txt) < 80 else "paragrafo"
                doc.blocos.append(Bloco(tipo=tipo, spans=_parse_inline(txt)))
            paragrafo_buffer.clear()

    for raw in linhas:
        l = raw.strip()
        if not l:
            flush_paragrafo()
            continue
        if _is_separator(l):
            flush_paragrafo()
            continue

        nivel, txt = _strip_md_heading(l)
        if nivel:
            flush_paragrafo()
            doc.blocos.append(
                Bloco(tipo="heading", nivel=nivel, spans=_parse_inline(txt))
            )
            continue

        bullet = _strip_bullet(l)
        if bullet is not None:
            flush_paragrafo()
            doc.blocos.append(Bloco(tipo="bullet", spans=_parse_inline(bullet)))
            continue

        # Tudo mais é parágrafo (concatena linhas)
        paragrafo_buffer.append(l)

    flush_paragrafo()

    # Renormaliza níveis de heading: o menor nível usado vira nivel=1.
    # (Quando a LLM usa `# Titulo` + `## Section`, depois de consumir o título
    # como cabeçalho, queremos que `##` no corpo apareça como Heading 1.)
    body_h_levels = [b.nivel for b in doc.blocos if b.tipo == "heading"]
    if body_h_levels:
        base = min(body_h_levels)
        for b in doc.blocos:
            if b.tipo == "heading":
                b.nivel = max(1, b.nivel - base + 1)

    return doc


# --------------------------------------------------------------------------- #
# Paleta sóbria (corporativa cinza-escuro)
# --------------------------------------------------------------------------- #
PALETA = {
    "texto": "#1A1A1A",
    "secundario": "#4A4A4A",
    "sutil": "#7A7A7A",
    "regua": "#B5B5B5",
    "fundo_meta": "#F2F2F2",
    "accent": "#2C2C2C",  # bem escuro, quase preto, pra headings
}


# --------------------------------------------------------------------------- #
# PDF (ReportLab)
# --------------------------------------------------------------------------- #
def render_pdf(
    doc_ir: Documento,
    mapa_png: str | Path | None,
    dossie: dict,
    semana: dict,
    out_pdf: str | Path,
) -> None:
    """Renderiza o RELINT em PDF estilo sóbrio corporativo."""
    from reportlab.lib.colors import HexColor
    from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        BaseDocTemplate,
        Frame,
        Image,
        PageTemplate,
        Paragraph,
        Spacer,
        Table,
        TableStyle,
        KeepTogether,
    )

    out_pdf = Path(out_pdf)

    # ---- Estilos ----
    base = getSampleStyleSheet()
    cor_texto = HexColor(PALETA["texto"])
    cor_sec = HexColor(PALETA["secundario"])
    cor_sutil = HexColor(PALETA["sutil"])
    cor_regua = HexColor(PALETA["regua"])
    cor_meta = HexColor(PALETA["fundo_meta"])
    cor_accent = HexColor(PALETA["accent"])

    s_titulo = ParagraphStyle(
        "Titulo",
        parent=base["Normal"],
        fontName="Helvetica-Bold",
        fontSize=20,
        leading=24,
        textColor=cor_accent,
        spaceAfter=4,
    )
    s_subsidio = ParagraphStyle(
        "Subsidio",
        parent=base["Normal"],
        fontName="Helvetica",
        fontSize=10,
        leading=13,
        textColor=cor_sutil,
        spaceAfter=10,
    )
    s_zona = ParagraphStyle(
        "ZonaTitulo",
        parent=base["Normal"],
        fontName="Helvetica-Bold",
        fontSize=14,
        leading=18,
        textColor=cor_texto,
        spaceAfter=4,
    )
    s_meta = ParagraphStyle(
        "Meta",
        parent=base["Normal"],
        fontName="Helvetica",
        fontSize=8.5,
        leading=11,
        textColor=cor_sec,
    )
    s_h1 = ParagraphStyle(
        "H1",
        parent=base["Normal"],
        fontName="Helvetica-Bold",
        fontSize=12,
        leading=16,
        textColor=cor_accent,
        spaceBefore=14,
        spaceAfter=4,
        keepWithNext=1,
    )
    s_h2 = ParagraphStyle(
        "H2",
        parent=s_h1,
        fontSize=11,
        spaceBefore=10,
    )
    s_corpo = ParagraphStyle(
        "Corpo",
        parent=base["Normal"],
        fontName="Helvetica",
        fontSize=10,
        leading=14,
        textColor=cor_texto,
        alignment=TA_JUSTIFY,
        spaceAfter=6,
    )
    s_lista_intro = ParagraphStyle(
        "ListaIntro",
        parent=s_corpo,
        spaceAfter=2,
    )
    s_bullet = ParagraphStyle(
        "Bullet",
        parent=s_corpo,
        leftIndent=16,
        firstLineIndent=-10,
        bulletIndent=2,
        spaceAfter=3,
        leading=13,
    )
    s_caption = ParagraphStyle(
        "Caption",
        parent=base["Normal"],
        fontName="Helvetica-Oblique",
        fontSize=8.5,
        leading=11,
        textColor=cor_sutil,
        alignment=TA_LEFT,
        spaceAfter=10,
    )

    # ---- Frame e template ----
    margem = 2.0 * cm
    cabecalho_h = 1.6 * cm
    rodape_h = 1.0 * cm
    largura, altura = A4

    frame = Frame(
        margem,
        margem + rodape_h,
        largura - 2 * margem,
        altura - 2 * margem - cabecalho_h - rodape_h,
        leftPadding=0,
        rightPadding=0,
        topPadding=0,
        bottomPadding=0,
    )

    cabecalho_texto = "COMPSTAT RIO  ·  RELATÓRIO DE INTELIGÊNCIA DE ÁREA"
    documento_id = (
        f"RA-{dossie.get('prioridade', 0):03d} · {semana.get('iso_ano')}-S{semana.get('iso_sem'):02d}"
    )

    def desenha_chrome(canv, _doc):
        canv.saveState()
        # Cabeçalho
        canv.setFont("Helvetica-Bold", 8)
        canv.setFillColor(cor_sec)
        canv.drawString(margem, altura - margem - 4, cabecalho_texto)
        canv.setFont("Helvetica", 8)
        canv.setFillColor(cor_sutil)
        canv.drawRightString(
            largura - margem,
            altura - margem - 4,
            documento_id,
        )
        canv.setStrokeColor(cor_regua)
        canv.setLineWidth(0.4)
        canv.line(
            margem,
            altura - margem - 10,
            largura - margem,
            altura - margem - 10,
        )
        # Rodapé
        canv.setFont("Helvetica", 8)
        canv.setFillColor(cor_sutil)
        canv.drawString(margem, margem - 4, dossie.get("local_rotulo", ""))
        canv.drawRightString(
            largura - margem, margem - 4, f"Página {canv.getPageNumber()}"
        )
        canv.restoreState()

    pdf_doc = BaseDocTemplate(
        str(out_pdf),
        pagesize=A4,
        leftMargin=margem,
        rightMargin=margem,
        topMargin=margem + cabecalho_h,
        bottomMargin=margem + rodape_h,
        title=f"RELINT — {dossie.get('local_rotulo', '')}",
        author="CompStat Rio · Group 14",
    )
    pdf_doc.addPageTemplates(
        [PageTemplate(id="default", frames=[frame], onPage=desenha_chrome)]
    )

    # ---- Helpers para spans → markup do ReportLab ----
    def spans_to_markup(spans: list[Span]) -> str:
        """ReportLab Paragraph aceita HTML-like. Escapa caracteres perigosos."""
        partes = []
        for sp in spans:
            t = (
                sp.text.replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
            )
            if sp.bold and sp.italic:
                t = f"<b><i>{t}</i></b>"
            elif sp.bold:
                t = f"<b>{t}</b>"
            elif sp.italic:
                t = f"<i>{t}</i>"
            partes.append(t)
        return "".join(partes)

    story: list = []

    # ---- Capa ----
    titulo = doc_ir.titulo or "RELATÓRIO DE INTELIGÊNCIA DE ÁREA"
    story.append(Paragraph(titulo, s_titulo))
    if doc_ir.subsidio:
        story.append(Paragraph(doc_ir.subsidio, s_subsidio))
    zona_label = doc_ir.subtitulo_zona or dossie.get("local_rotulo", "")
    if zona_label:
        story.append(Paragraph(zona_label, s_zona))

    # Bloco de metadados
    prio = dossie.get("prioridade", "?")
    n_total = dossie.get("n_zonas_sem")
    prio_str = f"#{prio}" + (f" de {n_total}" if n_total else "")
    meta_rows = [
        [
            Paragraph(f"<b>Semana</b><br/>{semana.get('iso_ano')}-S{semana.get('iso_sem'):02d}", s_meta),
            Paragraph(f"<b>Segunda-feira</b><br/>{semana.get('segunda_feira', '')}", s_meta),
            Paragraph(f"<b>Prioridade</b><br/>{prio_str}", s_meta),
            Paragraph(f"<b>Agentes alocados</b><br/>{dossie.get('agentes_alocados', '?')}", s_meta),
            Paragraph(f"<b>Área</b><br/>{dossie.get('area_km2', '?')} km² · {dossie.get('n_celulas', '?')} células", s_meta),
        ]
    ]
    largura_util = largura - 2 * margem
    cols = [largura_util / 5] * 5
    meta_tab = Table(meta_rows, colWidths=cols, hAlign="LEFT")
    meta_tab.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), cor_meta),
                ("TEXTCOLOR", (0, 0), (-1, -1), cor_sec),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                ("LINEBELOW", (0, 0), (-1, -1), 0.4, cor_regua),
                ("LINEABOVE", (0, 0), (-1, -1), 0.4, cor_regua),
            ]
        )
    )
    story.append(meta_tab)
    story.append(Spacer(1, 0.5 * cm))

    # ---- Mapa ----
    if mapa_png and Path(mapa_png).exists():
        from PIL import Image as PILImage

        with PILImage.open(mapa_png) as im:
            iw, ih = im.size
        max_w = largura_util
        max_h = 9 * cm
        ratio = min(max_w / iw, max_h / (ih * 72 / 96))
        w = iw * ratio
        h = ih * ratio
        img = Image(str(mapa_png), width=w, height=h)
        img.hAlign = "CENTER"
        story.append(img)
        story.append(
            Paragraph(
                f"<b>Mapa.</b> Polígono da zona com ocorrências (vermelho), denúncias do "
                f"Disque (laranja) e fatores urbanos (azul) no entorno. "
                f"Centróide: {dossie.get('centroide', {}).get('lat')}, "
                f"{dossie.get('centroide', {}).get('lon')}.",
                s_caption,
            )
        )

    # ---- Corpo ----
    for bloco in doc_ir.blocos:
        markup = spans_to_markup(bloco.spans)
        if bloco.tipo == "heading":
            estilo = s_h1 if bloco.nivel <= 1 else s_h2
            story.append(Paragraph(markup, estilo))
        elif bloco.tipo == "bullet":
            story.append(Paragraph(markup, s_bullet, bulletText="\u2022"))
        elif bloco.tipo == "lista_intro":
            story.append(Paragraph(markup, s_lista_intro))
        else:
            story.append(Paragraph(markup, s_corpo))

    pdf_doc.build(story)


# --------------------------------------------------------------------------- #
# DOCX (python-docx) — enxuto, focado em editabilidade
# --------------------------------------------------------------------------- #
def render_docx(
    doc_ir: Documento,
    mapa_png: str | Path | None,
    dossie: dict,
    semana: dict,
    out_docx: str | Path,
) -> None:
    """Renderiza o RELINT em DOCX usando estilos nativos do Word.

    Filosofia: simples, enxuto, sem fancy banners — facilita edição manual
    no Word/LibreOffice. Usa Title / Heading 1 / Heading 2 / List Bullet /
    Caption como estilos primários.
    """
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    out_docx = Path(out_docx)
    doc = Document()

    # Margens razoáveis
    for sec in doc.sections:
        sec.top_margin = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(2.2)
        sec.right_margin = Cm(2.2)

    # Tweaks nos estilos default
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    def tweak_heading(style_name: str, size_pt: int):
        st = styles[style_name]
        st.font.name = "Calibri"
        st.font.size = Pt(size_pt)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0x2C, 0x2C, 0x2C)

    tweak_heading("Title", 22)
    tweak_heading("Heading 1", 13)
    tweak_heading("Heading 2", 11)

    # Parágrafo helper: adiciona spans com bold/italic preservados
    def add_paragraph_with_spans(spans: list[Span], style: str | None = None):
        p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
        for sp in spans:
            r = p.add_run(sp.text)
            if sp.bold:
                r.bold = True
            if sp.italic:
                r.italic = True
        return p

    # ---- Cabeçalho do documento ----
    titulo = doc_ir.titulo or "RELATÓRIO DE INTELIGÊNCIA DE ÁREA"
    p = doc.add_paragraph(titulo, style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if doc_ir.subsidio:
        p = doc.add_paragraph()
        r = p.add_run(doc_ir.subsidio)
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x7A, 0x7A, 0x7A)

    if doc_ir.subtitulo_zona:
        p = doc.add_paragraph()
        r = p.add_run(doc_ir.subtitulo_zona)
        r.bold = True
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    # Linha de metadados (texto compacto, single-line, fonte menor)
    prio = dossie.get("prioridade", "?")
    n_total = dossie.get("n_zonas_sem")
    prio_str = f"#{prio}" + (f" de {n_total}" if n_total else "")
    meta = (
        f"Semana {semana.get('iso_ano')}-S{semana.get('iso_sem'):02d} · "
        f"Segunda-feira {semana.get('segunda_feira', '')} · "
        f"Prioridade {prio_str} · "
        f"{dossie.get('agentes_alocados', '?')} agentes · "
        f"{dossie.get('area_km2', '?')} km² · "
        f"{dossie.get('n_celulas', '?')} células"
    )
    p = doc.add_paragraph()
    r = p.add_run(meta)
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

    # Mapa
    if mapa_png and Path(mapa_png).exists():
        doc.add_picture(str(mapa_png), width=Cm(16.5))
        cap = doc.add_paragraph()
        cr = cap.add_run(
            f"Mapa: polígono da zona com ocorrências (vermelho), denúncias do "
            f"Disque (laranja) e fatores urbanos (azul) no entorno. "
            f"Centróide: {dossie.get('centroide', {}).get('lat')}, "
            f"{dossie.get('centroide', {}).get('lon')}."
        )
        cr.italic = True
        cr.font.size = Pt(9)
        cr.font.color.rgb = RGBColor(0x7A, 0x7A, 0x7A)

    # ---- Corpo ----
    for bloco in doc_ir.blocos:
        if bloco.tipo == "heading":
            estilo = "Heading 1" if bloco.nivel <= 1 else "Heading 2"
            # Heading não usa runs com bold individual — concatena tudo
            texto = "".join(sp.text for sp in bloco.spans)
            doc.add_paragraph(texto, style=estilo)
        elif bloco.tipo == "bullet":
            add_paragraph_with_spans(bloco.spans, style="List Bullet")
        else:
            add_paragraph_with_spans(bloco.spans)

    doc.save(str(out_docx))


# --------------------------------------------------------------------------- #
# CLI de smoke-test
# --------------------------------------------------------------------------- #
def _smoke_test():
    """Re-renderiza o RA_001_Centro a partir do texto já existente no docx
    antigo, sem chamar a API. Bom pra iterar no estilo."""
    import argparse
    from docx import Document as _D

    BASE = Path(__file__).resolve().parent
    # O smoke test extrai o texto bruto da LLM de um docx existente (sem chamar
    # a API). Usamos o backup do docx mal-formatado como fonte para que o smoke
    # test reproduza fielmente o input original (com markdown literal `# X`,
    # `**X**`, `---`).
    DEFAULT_DOCX = BASE / "distribuicao_fm" / "relatorios_ia" / "RA_001_Centro_BAD_BACKUP.docx"
    DEFAULT_MAPA = BASE / "distribuicao_fm" / "relatorios_ia" / "RA_001_Centro_mapa.png"
    DEFAULT_DOSSIE_JSON = BASE / "distribuicao_fm" / "relatorio_zonas_compacto.json"

    p = argparse.ArgumentParser(description=__doc__)
    p.add_argument("--docx-fonte", type=Path, default=DEFAULT_DOCX,
                   help="DOCX antigo de onde extrair o texto bruto do LLM")
    p.add_argument("--mapa", type=Path, default=DEFAULT_MAPA)
    p.add_argument("--dossie-json", type=Path, default=DEFAULT_DOSSIE_JSON)
    p.add_argument("--zona-id", type=int, default=1)
    p.add_argument("--out-pdf", type=Path,
                   default=BASE / "distribuicao_fm" / "relatorios_ia" / "RA_001_Centro.pdf")
    p.add_argument("--out-docx", type=Path,
                   default=BASE / "distribuicao_fm" / "relatorios_ia" / "RA_001_Centro_NEW.docx")
    args = p.parse_args()

    # Extrai o texto bruto do docx antigo. Pula até a primeira linha que parece
    # o início do RELINT (a `# RELATORIO ...`). Para parágrafos com estilo
    # List Bullet, prefixa com "- " (o p.text não inclui o marcador).
    d = _D(str(args.docx_fonte))
    linhas: list[str] = []
    inicio = False
    for p in d.paragraphs:
        t = p.text.strip()
        if not inicio:
            if t.startswith("#"):
                inicio = True
            else:
                continue
        if not t:
            linhas.append("")
            continue
        if "Bullet" in p.style.name or "List" in p.style.name:
            linhas.append(f"- {t}")
        else:
            linhas.append(t)
    texto_llm = "\n".join(linhas)

    import json
    payload = json.loads(args.dossie_json.read_text(encoding="utf-8"))
    semana = payload["semana"]
    dossie = next((z for z in payload["zonas"] if z["zona_id"] == args.zona_id), None)
    if dossie is None:
        raise SystemExit(f"zona_id={args.zona_id} não encontrada em {args.dossie_json}")
    dossie["n_zonas_sem"] = payload["parametros"]["n_zonas"]

    print(f"[parse] {len(texto_llm.splitlines())} linhas de input")
    doc_ir = parse_relint(texto_llm)
    print(f"[parse] título='{doc_ir.titulo}'")
    print(f"[parse] subsídio='{doc_ir.subsidio}'")
    print(f"[parse] subtítulo='{doc_ir.subtitulo_zona}'")
    print(f"[parse] {len(doc_ir.blocos)} blocos no corpo")

    render_pdf(doc_ir, args.mapa, dossie, semana, args.out_pdf)
    print(f"[ok]    PDF  -> {args.out_pdf}  ({args.out_pdf.stat().st_size/1024:.1f} KB)")
    render_docx(doc_ir, args.mapa, dossie, semana, args.out_docx)
    print(f"[ok]    DOCX -> {args.out_docx}  ({args.out_docx.stat().st_size/1024:.1f} KB)")


if __name__ == "__main__":
    _smoke_test()
